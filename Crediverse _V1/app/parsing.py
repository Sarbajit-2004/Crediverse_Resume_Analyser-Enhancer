from __future__ import annotations
from pypdf import PdfReader
from docx import Document

def extract_text_from_pdf(path: str) -> str:
    r = PdfReader(path)
    return "\n".join([(p.extract_text() or "") for p in r.pages]).strip()

def extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    lines = [p.text for p in doc.paragraphs if p.text]
    for t in doc.tables:
        for row in t.rows:
            lines.append(" | ".join(c.text for c in row.cells))
    return "\n".join(lines).strip()
