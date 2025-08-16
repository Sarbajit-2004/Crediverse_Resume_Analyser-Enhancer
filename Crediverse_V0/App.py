import os, io, uuid, base64, time, datetime, re
import streamlit as st
import pandas as pd
from PIL import Image
from streamlit_tags import st_tags
from pypdf import PdfReader
from docx import Document
import pymysql
import plotly.express as px

# ---------------------------
# PAGE CONFIG
# ---------------------------
st.set_page_config(page_title="AI Resume Analyzer", page_icon='./Logo/logo2.png')

UPLOAD_DIR = "./Uploaded_Resumes"
LOGO_PATH = "./Logo/logo2.png"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# ---------------------------
# COURSES
# ---------------------------
from Courses import (
    ds_course, web_course, android_course, ios_course, uiux_course,
    resume_videos, interview_videos
)

# ---------------------------
# DB (use Streamlit secrets or env)
# .streamlit/secrets.toml:
# [mysql]
# host="localhost"
# user="root"
# password="YOUR_PASSWORD"
# db="cv"
# ---------------------------
def get_db_conn(db: str | None = None):
    cfg = st.secrets.get("mysql", {})
    host = cfg.get("host", os.getenv("MYSQL_HOST", "localhost"))
    user = cfg.get("user", os.getenv("MYSQL_USER", "root"))
    password = cfg.get("password", os.getenv("MYSQL_PASSWORD", ""))
    if db is None:
        return pymysql.connect(host=host, user=user, password=password, autocommit=True)
    return pymysql.connect(host=host, user=user, password=password, database=db, autocommit=True)

def init_db():
    # create DB if missing
    conn = get_db_conn(None)
    cur = conn.cursor()
    cur.execute("CREATE DATABASE IF NOT EXISTS cv CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci;")
    cur.close(); conn.close()

    # create table if missing
    conn = get_db_conn("cv")
    cur = conn.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS user_data(
            ID INT NOT NULL AUTO_INCREMENT PRIMARY KEY,
            Name VARCHAR(200),
            Email_ID VARCHAR(200),
            Resume_Score INT,
            Timestamp VARCHAR(50),
            Page_no INT,
            Predicted_Field VARCHAR(100),
            User_level VARCHAR(50),
            Actual_skills TEXT,
            Recommended_skills TEXT,
            Recommended_courses TEXT
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
    """)
    cur.close(); conn.close()

def insert_data(name,email,res_score,timestamp,no_of_pages,reco_field,cand_level,skills,recommended_skills,courses):
    conn = get_db_conn("cv")
    cur = conn.cursor()
    sql = """
    INSERT INTO user_data
    (Name,Email_ID,Resume_Score,Timestamp,Page_no,Predicted_Field,User_level,Actual_skills,Recommended_skills,Recommended_courses)
    VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
    """
    cur.execute(sql, (
        name or "", email or "", int(res_score), timestamp, int(no_of_pages),
        reco_field, cand_level, skills, recommended_skills, courses
    ))
    cur.close(); conn.close()

# ---------------------------
# HELPERS
# ---------------------------
def show_pdf(file_path: str):
    with open(file_path, "rb") as f:
        base64_pdf = base64.b64encode(f.read()).decode('utf-8')
    st.markdown(
        f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="700" height="1000" type="application/pdf"></iframe>',
        unsafe_allow_html=True,
    )

def extract_text_from_pdf(path: str) -> str:
    reader = PdfReader(path)
    parts = []
    for p in reader.pages:
        parts.append(p.extract_text() or "")
    return "\n".join(parts).strip()

def extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    lines = [p.text for p in doc.paragraphs if p.text]
    for t in doc.tables:
        for r in t.rows:
            lines.append(" | ".join(c.text for c in r.cells))
    return "\n".join(lines).strip()

def detect_basic_fields(text: str):
    email = None; phone = None
    m = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text);  email = m.group(0) if m else None
    m = re.search(r'(\+?\d[\d\s\-\(\)]{8,}\d)', text); phone = m.group(0) if m else None
    first_line = (text.splitlines() or [""])[0].strip()
    name_guess = first_line if 3 <= len(first_line) <= 60 else None
    return {"name": name_guess, "email": email, "mobile_number": phone}

def course_recommender(course_list):
    st.subheader("**Courses & Certificates Recommendations 🎓**")
    rec_course = []
    no_of_reco = st.slider('Choose Number of Course Recommendations:', 1, 10, 5)
    import random; pool = list(course_list); random.shuffle(pool)
    for i, (title, link) in enumerate(pool[:no_of_reco], start=1):
        st.markdown(f"({i}) [{title}]({link})")
        rec_course.append(title)
    return rec_course

def ats_keywords(text: str):
    # very small stopword list to keep it simple (no NLTK)
    stop = set("""the a an and or of in on for with to from by as at is are be this that those these your you we our their i me my""".split())
    words = re.findall(r"[A-Za-z][A-Za-z0-9+\-#\.]*", text.lower())
    return [w for w in words if w not in stop and len(w) > 2]

# ---------------------------
# MAIN APP
# ---------------------------
def run():
    init_db()

    if os.path.exists(LOGO_PATH):
        # Option A (recommended): let Streamlit pick the natural size
        st.image(Image.open(LOGO_PATH))

    st.title("AI Resume Analyser")
    st.sidebar.markdown("# Choose User")
    choice = st.sidebar.selectbox("Choose among the given options:", ["User", "Admin"])
    st.sidebar.markdown('[©Developed by Sarbajit](https://www.linkedin.com/public-profile/settings?trk=d_flagship3_profile_self_view_public_profile)', unsafe_allow_html=True)

    if choice == "User":
        st.markdown("**Upload your resume (PDF or DOCX), and get smart recommendations**")
        up = st.file_uploader("Choose your Resume", type=["pdf", "docx"])
        if up is not None:
            ext = os.path.splitext(up.name)[1].lower()
            safe_name = f"{uuid.uuid4().hex}{ext}"
            save_path = os.path.join(UPLOAD_DIR, safe_name)
            with open(save_path, "wb") as f:
                f.write(up.getbuffer())

            # Preview PDF only
            if ext == ".pdf":
                show_pdf(save_path)

            # Extract text + pages
            if ext == ".pdf":
                resume_text = extract_text_from_pdf(save_path)
                pages = len(PdfReader(save_path).pages)
            else:
                resume_text = extract_text_from_docx(save_path)
                pages = 1  # docx page count isn't trivial; treat as 1

            fields = detect_basic_fields(resume_text)

            st.header("**Resume Analysis**")
            st.success(f"Hello {fields.get('name') or 'Candidate'}")
            st.subheader("**Your Basic info**")
            st.text(f"Name: {fields.get('name') or 'N/A'}")
            st.text(f"Email: {fields.get('email') or 'N/A'}")
            st.text(f"Contact: {fields.get('mobile_number') or 'N/A'}")
            st.text(f"Resume pages: {pages}")

            # Current skills (user editable)
            st.subheader("**Your Current Skills**")
            keywords = st_tags(
                label='### Add/Confirm your skills',
                text='Type and press enter',
                value=[],
                key='skills_tags'
            )

            # Field detection via skills
            ds_keyword = ['tensorflow','keras','pytorch','machine learning','deep learning','flask','streamlit']
            web_keyword = ['react','django','node js','react js','php','laravel','magento','wordpress','javascript','angular js','c#','flask']
            android_keyword = ['android','android development','flutter','kotlin','xml','kivy']
            ios_keyword = ['ios','ios development','swift','cocoa','cocoa touch','xcode']
            uiux_keyword = ['ux','adobe xd','figma','zeplin','balsamiq','ui','prototyping','wireframes','photoshop','illustrator','after effects','premiere pro','indesign','user research']

            def any_in(skills, pool):
                return any(s.lower() in pool for s in skills)

            reco_field = ''
            recommended_skills = []
            rec_course = []

            if any_in(keywords, ds_keyword):
                reco_field = 'Data Science'
                st.success("**Our analysis says you are looking for Data Science Jobs.**")
                recommended_skills = ['Data Visualization','Predictive Analysis','Statistical Modeling','Data Mining',
                                      'Clustering & Classification','Data Analytics','Quantitative Analysis','Web Scraping',
                                      'ML Algorithms','Keras','Pytorch','Probability','Scikit-learn','Tensorflow','Flask','Streamlit']
                st_tags(label='### Recommended skills for you.',
                        text='Recommended skills generated from System',
                        value=recommended_skills, key='reco1')
                rec_course = course_recommender(ds_course)

            elif any_in(keywords, web_keyword):
                reco_field = 'Web Development'
                st.success("**Our analysis says you are looking for Web Development Jobs.**")
                recommended_skills = ['React','Django','Node JS','React JS','PHP','Laravel','Magento','WordPress',
                                      'JavaScript','Angular','C#','Flask','SDK']
                st_tags(label='### Recommended skills for you.',
                        text='Recommended skills generated from System',
                        value=recommended_skills, key='reco2')
                rec_course = course_recommender(web_course)

            elif any_in(keywords, android_keyword):
                reco_field = 'Android Development'
                st.success("**Our analysis says you are looking for Android App Development Jobs.**")
                recommended_skills = ['Android','Flutter','Kotlin','XML','Java','Kivy','Git','SDK','SQLite']
                st_tags(label='### Recommended skills for you.',
                        text='Recommended skills generated from System',
                        value=recommended_skills, key='reco3')
                rec_course = course_recommender(android_course)

            elif any_in(keywords, ios_keyword):
                reco_field = 'iOS Development'
                st.success("**Our analysis says you are looking for iOS App Development Jobs.**")
                recommended_skills = ['Swift','Cocoa','Cocoa Touch','Xcode','Objective-C','SQLite','StoreKit','UIKit','AVFoundation','Auto Layout']
                st_tags(label='### Recommended skills for you.',
                        text='Recommended skills generated from System',
                        value=recommended_skills, key='reco4')
                rec_course = course_recommender(ios_course)

            elif any_in(keywords, uiux_keyword):
                reco_field = 'UI/UX Development'
                st.success("**Our analysis says you are looking for UI/UX Development Jobs.**")
                recommended_skills = ['User Experience','Adobe XD','Figma','Zeplin','Balsamiq','Prototyping','Wireframes',
                                      'Storyboards','Photoshop','Illustrator','After Effects','Premiere Pro','InDesign','User Research']
                st_tags(label='### Recommended skills for you.',
                        text='Recommended skills generated from System',
                        value=recommended_skills, key='reco5')
                rec_course = course_recommender(uiux_course)

            # ---------------------------
            # Resume scoring (keyword presence; case-insensitive)
            # ---------------------------
            st.subheader("**Resume Tips & Ideas💡**")
            resume_lower = resume_text.lower()
            score_items = [
                ("objective", "Add a clear career objective (optional but useful for freshers)."),
                ("declaration", "Add a short declaration/signature block if desired."),
                ("hobbies", "Consider a small hobbies/interests line to show personality."),
                ("achievements", "Add achievements with measurable outcomes."),
                ("projects", "List projects with tools, responsibility, and impact.")
            ]
            resume_score = 0
            for key, tip in score_items:
                if key in resume_lower:
                    resume_score += 20
                    st.markdown(f"<h5 style='color:#1ed760'>[+] {key.title()} found</h5>", unsafe_allow_html=True)
                else:
                    st.markdown(f"<h5 style='color:#333'>[-] {tip}</h5>", unsafe_allow_html=True)

            st.subheader("**Resume Score📝**")
            st.progress(resume_score)
            st.success(f"** Your Resume Writing Score: {resume_score} **")
            st.info("Heuristic score based on presence of key sections/keywords.")

            # ---------------------------
            # Simple ATS keyword match vs JD
            # ---------------------------
            st.subheader("**ATS Keyword Match (paste a Job Description)**")
            jd = st.text_area("Paste JD here (optional)")
            if jd.strip():
                res_keys = set(ats_keywords(resume_text))
                jd_keys = set(ats_keywords(jd))
                present = sorted(jd_keys & res_keys)
                missing = sorted(jd_keys - res_keys)
                coverage = 0 if not jd_keys else round(100 * len(present) / len(jd_keys))
                st.write(f"**Coverage:** {coverage}%")
                st.write("**Present keywords:**", ", ".join(present) or "—")
                st.write("**Missing keywords:**", ", ".join(missing) or "—")

            # Insert to DB
            ts = time.time()
            timestamp = f"{datetime.datetime.fromtimestamp(ts):%Y-%m-%d_%H:%M:%S}"
            cand_level = "Fresher" if pages == 1 else ("Intermediate" if pages == 2 else "Experienced")

            insert_data(
                fields.get("name"), fields.get("email"),
                resume_score, timestamp, pages, reco_field, cand_level,
                ", ".join(keywords), ", ".join(recommended_skills),
                ", ".join(rec_course) if isinstance(rec_course, (list, tuple)) else str(rec_course),
            )

            # Videos
            st.header("**Bonus Video for Resume Writing Tips💡**")
            if resume_videos: st.video(resume_videos[0])

            st.header("**Bonus Video for Interview Tips💡**")
            if interview_videos: st.video(interview_videos[0])

    else:
        # Admin
        st.success('Welcome to Admin Side')
        ad_user = st.text_input("Username")
        ad_password = st.text_input("Password", type='password')
        if st.button('Login'):
            if ad_user == 'Sarbajit' and ad_password == 'Sarbajit123':
                st.success("Welcome Mr Sarbajit !")
                conn = get_db_conn("cv")
                df = pd.read_sql("SELECT * FROM user_data ORDER BY ID DESC", conn)
                conn.close()

                st.header("**User's Data**")
                st.dataframe(df)
                # CSV download
                csv = df.to_csv(index=False)
                b64 = base64.b64encode(csv.encode()).decode()
                st.markdown(f'<a href="data:file/csv;base64,{b64}" download="User_Data.csv">Download Report</a>', unsafe_allow_html=True)

                if not df.empty:
                    st.subheader("**Pie-Chart: Predicted Field Recommendation**")
                    fig1 = px.pie(df, names="Predicted_Field", title="Predicted Field by Skills")
                    st.plotly_chart(fig1, use_container_width=True)

                    st.subheader("**Pie-Chart: User Experience Level**")
                    fig2 = px.pie(df, names="User_level", title="User Level Distribution")
                    st.plotly_chart(fig2, use_container_width=True)
                else:
                    st.info("No records yet.")
            else:
                st.error("Wrong ID & Password Provided")

if __name__ == "__main__":
    run()
